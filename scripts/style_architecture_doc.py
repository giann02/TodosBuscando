from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "TodosBuscando_ArquitecturaBD.docx"

TABLE_WIDTH = 8400


def get_or_add(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_attr(element, name: str, value: str) -> None:
    element.set(qn(name), value)


def set_table_outer_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = get_or_add(tbl_pr, "w:tblW")
    set_attr(tbl_w, "w:type", "dxa")
    set_attr(tbl_w, "w:w", str(TABLE_WIDTH))

    borders = get_or_add(tbl_pr, "w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = get_or_add(borders, f"w:{edge}")
        set_attr(element, "w:val", "single")
        set_attr(element, "w:color", "auto")
        set_attr(element, "w:sz", "4")
        if element.get(qn("w:space")) is not None:
            set_attr(element, "w:space", "0")


def set_cell_props(cell, width: int, is_header: bool) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()

    tc_w = get_or_add(tc_pr, "w:tcW")
    set_attr(tc_w, "w:type", "dxa")
    set_attr(tc_w, "w:w", str(width))

    borders = get_or_add(tc_pr, "w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = get_or_add(borders, f"w:{edge}")
        set_attr(element, "w:val", "single")
        set_attr(element, "w:color", "CCCCCC")
        set_attr(element, "w:sz", "1")

    margins = get_or_add(tc_pr, "w:tcMar")
    for edge, value in (("top", "80"), ("left", "120"), ("bottom", "80"), ("right", "120")):
        element = get_or_add(margins, f"w:{edge}")
        set_attr(element, "w:type", "dxa")
        set_attr(element, "w:w", value)

    valign = get_or_add(tc_pr, "w:vAlign")
    set_attr(valign, "w:val", "center")

    shading = tc_pr.find(qn("w:shd"))
    if is_header:
        if shading is None:
            shading = OxmlElement("w:shd")
            tc_pr.append(shading)
        set_attr(shading, "w:fill", "C8DCF0")
        set_attr(shading, "w:val", "clear")
    elif shading is not None:
        tc_pr.remove(shading)


def table_widths(column_count: int) -> list[int]:
    if column_count == 4:
        return [1600, 1700, 1400, 3700]
    if column_count == 3:
        return [2200, 2400, 3800]
    width = TABLE_WIDTH // column_count
    widths = [width] * column_count
    widths[-1] += TABLE_WIDTH - sum(widths)
    return widths


def style_table(table) -> None:
    set_table_outer_borders(table)
    widths = table_widths(len(table.columns))

    for row_index, row in enumerate(table.rows):
        if row_index == 0:
            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:tblHeader")) is None:
                tr_pr.append(OxmlElement("w:tblHeader"))

        for cell_index, cell in enumerate(row.cells):
            is_header = row_index == 0
            set_cell_props(cell, widths[cell_index], is_header)

            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(11)
                    run.bold = True if is_header else False


def find_paragraph_style(doc: Document, prefix: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph.style
    return None


def style_section_8_text(doc: Document) -> None:
    heading_1_style = find_paragraph_style(doc, "7. Conclusiones")
    heading_2_style = find_paragraph_style(doc, "2.5. Resumen de Decisiones")

    in_section_8 = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("8. Modelo de Datos Detallado"):
            in_section_8 = True
        if not in_section_8:
            continue

        if text.startswith("8. Modelo de Datos Detallado"):
            if heading_1_style is not None:
                paragraph.style = heading_1_style
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = None
                run.bold = None
            continue

        if text.startswith("8.") and text[:4].count(".") >= 2:
            if heading_2_style is not None:
                paragraph.style = heading_2_style
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = None
                run.bold = None
            continue

        for run in paragraph.runs:
            run.font.name = "Arial"
            if text:
                run.font.size = Pt(12)


def main() -> None:
    doc = Document(DOCX_PATH)
    style_section_8_text(doc)

    # Tables 0 and 1 already belong to the original document. Tables 2 onward
    # were added for section 8 and should visually match table 2.5.
    for table in doc.tables[2:]:
        style_table(table)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
