from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "TodosBuscando_ArquitecturaBD.docx"
BACKUP_PATH = ROOT / "TodosBuscando_ArquitecturaBD.backup.docx"


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "BFBFBF")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)

    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)

    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 13)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"- {item}")
        run.font.size = Pt(10)


def add_code_like_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    if not BACKUP_PATH.exists():
        shutil.copyfile(DOCX_PATH, BACKUP_PATH)

    doc = Document(DOCX_PATH)

    doc.add_page_break()
    add_heading(doc, "8. Modelo de Datos Detallado", level=1)
    doc.add_paragraph(
        "Esta sección especifica con mayor precisión cómo se modelan las entidades principales del sistema. "
        "Aunque MongoDB no utiliza tablas en sentido relacional, se describen las colecciones como tablas lógicas "
        "para facilitar la lectura del modelo de datos y la defensa del trabajo práctico."
    )

    add_heading(doc, "8.1. Colección alertas", level=2)
    doc.add_paragraph(
        "La colección alertas representa cada caso de persona desaparecida. Es la entidad central del sistema: "
        "se crea de forma manual desde el panel de administración o de forma automática a partir de la API externa. "
        "Cada documento concentra los datos descriptivos del caso, su estado operativo y su ubicación geográfica."
    )
    add_table(
        doc,
        ["Campo", "Tipo", "Obligatorio", "Descripción"],
        [
            ["_id / id", "ObjectId / String", "Sí", "Identificador único interno generado por MongoDB."],
            ["nombre", "String", "Sí", "Nombre de la persona desaparecida. Se utiliza también para búsqueda textual y detección de duplicados."],
            ["edad", "Integer", "Sí", "Edad de la persona al momento de la alerta."],
            ["descripcion", "String", "Sí", "Descripción física, vestimenta u otra información útil para reconocer a la persona."],
            ["fotoUrl", "String", "No", "Ruta local o URL de la imagen asociada a la alerta."],
            ["ultimaUbicacionConocida", "String", "Sí", "Dirección textual o referencia del último lugar donde fue vista la persona."],
            ["ubicacion", "GeoJSON Point", "No", "Coordenadas geográficas en formato longitud/latitud. Permite búsquedas geoespaciales."],
            ["estado", "String", "Sí", "Estado operativo de la alerta. Valores esperados: ACTIVA o RESUELTA."],
            ["vecinosNotificados", "Integer", "Sí", "Cantidad de usuarios cercanos que fueron notificados al crear la alerta."],
            ["creadoEn", "DateTime", "Sí", "Fecha y hora de creación o importación de la alerta."],
            ["origen", "String", "Sí", "Indica si la alerta fue creada MANUALMENTE o importada desde API_EXTERNA."],
            ["idExterno", "String", "No", "Identificador proveniente de la API externa. Se usa para evitar duplicados en la sincronización automática."],
        ],
    )
    add_table(
        doc,
        ["Índice", "Campos", "Uso"],
        [
            ["Índice primario", "_id", "Consultar una alerta puntual por identificador."],
            ["Índice por estado", "estado", "Listar rápidamente las alertas activas que se muestran en la web y el panel."],
            ["Índice geoespacial 2dsphere", "ubicacion", "Ubicar alertas en mapas y habilitar futuras consultas por cercanía."],
            ["Índice de texto", "nombre", "Detectar posibles duplicados al cargar una nueva alerta manual."],
            ["Índice único sparse", "idExterno", "Evitar importar dos veces el mismo caso desde la API externa."],
        ],
    )
    doc.add_paragraph("Ejemplo simplificado de documento:")
    add_code_like_paragraph(
        doc,
        "{ id: '665...', nombre: 'Juan Perez', edad: 15, estado: 'ACTIVA', "
        "origen: 'MANUAL', ubicacion: { type: 'Point', coordinates: [-58.3816, -34.6037] }, "
        "vecinosNotificados: 12, creadoEn: '2026-06-12T10:30:00' }",
    )

    add_heading(doc, "8.2. Colección usuarios", level=2)
    doc.add_paragraph(
        "La colección usuarios almacena a los vecinos registrados para recibir alertas. Su dato más importante "
        "es la ubicación, porque el sistema la utiliza para decidir a quién notificar cuando aparece una alerta cerca."
    )
    add_table(
        doc,
        ["Campo", "Tipo", "Obligatorio", "Descripción"],
        [
            ["_id / id", "ObjectId / String", "Sí", "Identificador único interno del usuario."],
            ["nombre", "String", "Sí", "Nombre del vecino registrado."],
            ["email", "String", "Sí", "Correo usado para enviar notificaciones. Debe ser único."],
            ["ubicacion", "GeoJSON Point", "Sí", "Punto geográfico donde el usuario desea recibir alertas de cercanía."],
            ["fechaRegistro", "DateTime", "Sí", "Momento en que el usuario se registró en el sistema."],
        ],
    )
    add_table(
        doc,
        ["Índice", "Campos", "Uso"],
        [
            ["Índice primario", "_id", "Actualizar o consultar un usuario específico."],
            ["Índice único", "email", "Evitar registros duplicados con el mismo correo electrónico."],
            ["Índice geoespacial 2dsphere", "ubicacion", "Buscar vecinos dentro de un radio de 2 km desde la alerta."],
        ],
    )
    doc.add_paragraph("Consulta principal asociada:")
    add_code_like_paragraph(
        doc,
        "findByUbicacionNear(Point alerta, Distance 2km) -> usuarios cercanos a notificar",
    )

    add_heading(doc, "8.3. Colección reportes", level=2)
    doc.add_paragraph(
        "La colección reportes registra cada avistamiento enviado por la comunidad. Un reporte pertenece a una "
        "alerta y puede incluir ubicación GPS, ubicación elegida manualmente o no incluir ubicación si la persona "
        "que reporta no puede precisarla."
    )
    add_table(
        doc,
        ["Campo", "Tipo", "Obligatorio", "Descripción"],
        [
            ["_id / id", "ObjectId / String", "Sí", "Identificador único del reporte en MongoDB."],
            ["alertaId", "String", "Sí", "Referencia lógica a la alerta sobre la que se está reportando el avistamiento."],
            ["descripcion", "String", "Sí", "Detalle libre del avistamiento informado."],
            ["reportadoPor", "String", "No", "Nombre de la persona que reporta. Puede ser nulo si el reporte es anónimo."],
            ["ubicacion", "GeoJSON Point", "No", "Coordenadas del avistamiento cuando están disponibles."],
            ["timestamp", "DateTime", "Sí", "Fecha y hora en que se registró el reporte."],
        ],
    )
    add_table(
        doc,
        ["Índice", "Campos", "Uso"],
        [
            ["Índice primario", "_id", "Consultar un reporte puntual."],
            ["Índice por alerta", "alertaId", "Listar todos los reportes asociados a una alerta y contar respuestas por caso."],
            ["Índice geoespacial opcional", "ubicacion", "Analizar distribución geográfica de avistamientos si crece el volumen de reportes."],
            ["Índice temporal opcional", "timestamp", "Ordenar reportes cronológicamente para análisis de trayectoria."],
        ],
    )

    add_heading(doc, "8.4. Nodos y relaciones en Neo4j", level=2)
    doc.add_paragraph(
        "Además de almacenarse en MongoDB, cada reporte con ubicación se replica como nodo en Neo4j. "
        "Este modelo no reemplaza a MongoDB: lo complementa para representar la secuencia temporal de avistamientos "
        "como un camino navegable."
    )
    add_table(
        doc,
        ["Elemento", "Propiedades", "Descripción"],
        [
            ["Nodo Reporte", "mongoId, alertaId, descripcion, lat, lng, timestamp", "Representa un avistamiento geolocalizado. mongoId mantiene la trazabilidad con MongoDB."],
            ["Relación SIGUIENTE", "Sin propiedades obligatorias", "Conecta un reporte con el siguiente reporte cronológico de la misma alerta."],
        ],
    )
    doc.add_paragraph("Patrón del grafo:")
    add_code_like_paragraph(
        doc,
        "(:Reporte {alertaId}) -[:SIGUIENTE]-> (:Reporte {alertaId}) -[:SIGUIENTE]-> (:Reporte {alertaId})",
    )
    add_bullets(
        doc,
        [
            "MongoDB conserva el documento completo del reporte.",
            "Neo4j conserva la información mínima necesaria para recorrer la trayectoria.",
            "La relación SIGUIENTE se crea únicamente entre reportes de la misma alerta.",
            "El orden se define por timestamp, por lo que la cadena representa una secuencia temporal.",
        ],
    )

    add_heading(doc, "8.5. Claves utilizadas en Redis", level=2)
    doc.add_paragraph(
        "Redis se utiliza como almacenamiento clave-valor temporal. No contiene la fuente de verdad del sistema, "
        "sino copias de lectura rápida que pueden eliminarse y reconstruirse desde MongoDB."
    )
    add_table(
        doc,
        ["Clave lógica", "Valor", "TTL", "Cuándo se invalida"],
        [
            ["alertas:activas", "Lista serializada de alertas con estado ACTIVA", "300 segundos", "Al crear una alerta nueva o resolver una alerta existente."],
            ["dedupe:api:<idExterno>", "Marca de caso importado desde API externa", "Según política de sincronización", "Puede expirar o mantenerse para evitar reprocesamientos repetidos."],
        ],
    )
    doc.add_paragraph(
        "El uso de Redis mejora el rendimiento de las lecturas repetidas, pero si Redis falla el sistema continúa "
        "funcionando porque consulta MongoDB directamente."
    )

    add_heading(doc, "8.6. Relaciones lógicas entre entidades", level=2)
    add_table(
        doc,
        ["Relación", "Cardinalidad", "Implementación"],
        [
            ["Alerta - Reporte", "1 a N", "Cada reporte guarda alertaId como referencia al id de la alerta."],
            ["Alerta - Usuario notificado", "N a N derivada", "No se guarda una tabla intermedia; se calcula por cercanía geográfica al momento de crear la alerta y se registra vecinosNotificados."],
            ["Reporte MongoDB - Reporte Neo4j", "1 a 1 opcional", "Si el reporte tiene ubicación, se crea un nodo en Neo4j con mongoId igual al id del reporte MongoDB."],
            ["Reporte Neo4j - Reporte Neo4j", "1 a 1 secuencial", "Cada nodo puede apuntar al siguiente avistamiento cronológico mediante SIGUIENTE."],
        ],
    )

    add_heading(doc, "8.7. Reglas de integridad y validaciones", level=2)
    add_bullets(
        doc,
        [
            "Una alerta manual debe tener nombre, edad, descripción y última ubicación conocida.",
            "Una alerta importada desde la API externa debe conservar idExterno para evitar duplicados.",
            "Un usuario no puede registrarse dos veces con el mismo email.",
            "La ubicación se almacena siempre como GeoJSON Point, respetando el orden longitud-latitud usado por MongoDB.",
            "Un reporte siempre debe pertenecer a una alerta existente mediante alertaId.",
            "Los reportes anónimos son válidos; por eso reportadoPor puede ser nulo.",
            "Solo los reportes con coordenadas generan nodos de trayectoria en Neo4j.",
            "Resolver una alerta cambia su estado y obliga a invalidar el caché de alertas activas.",
        ],
    )

    add_heading(doc, "8.8. Justificación del modelo elegido", level=2)
    doc.add_paragraph(
        "El modelo documental evita JOINs innecesarios en las operaciones más frecuentes. La pantalla principal "
        "necesita mostrar alertas completas, no combinar datos dispersos en varias tablas. Por eso cada alerta "
        "se guarda como documento autocontenido. En cambio, los reportes se separan porque crecen en cantidad "
        "a medida que la comunidad participa y porque se consultan por alerta."
    )
    doc.add_paragraph(
        "La decisión de no guardar una relación persistente entre alertas y usuarios notificados también es intencional: "
        "la pertenencia a la zona de una alerta se obtiene por consulta geoespacial. Guardar una tabla intermedia "
        "aumentaría la duplicación de datos y complicaría la actualización cuando un usuario cambia su ubicación."
    )
    doc.add_paragraph(
        "Neo4j se reserva para la parte del dominio que sí necesita relaciones explícitas: la trayectoria. "
        "Mientras MongoDB responde eficientemente consultas por documento, estado, texto y cercanía, Neo4j permite "
        "recorrer el camino de avistamientos como una secuencia conectada."
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
