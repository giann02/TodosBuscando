from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "TodosBuscando_ArquitecturaBD.docx"


def replace_paragraph(paragraph, text: str) -> None:
    for index, run in enumerate(paragraph.runs):
        run.text = text if index == 0 else ""


def set_run_style(paragraph) -> None:
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)


def main() -> None:
    doc = Document(DOCX_PATH)

    section_paragraph = None
    section_index = None
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "7.6. Relaciones lógicas entre entidades":
            section_paragraph = paragraph
            section_index = index
            replace_paragraph(paragraph, "7.6. Referencias y vínculos lógicos entre entidades")
            break

    if section_paragraph is None:
        raise RuntimeError("No se encontró la sección 7.6")

    clarification = (
        "En una base de datos NoSQL documental como MongoDB no existen relaciones en el sentido relacional clásico: "
        "no hay claves foráneas obligatorias, JOINs nativos ni integridad referencial automática entre colecciones. "
        "Lo que el sistema utiliza son referencias lógicas por identificador y consultas por criterio, por ejemplo "
        "guardar alertaId dentro de cada reporte para poder recuperar los reportes de una alerta. En Neo4j sí se "
        "modelan relaciones explícitas, porque su modelo es de grafos."
    )

    following_texts = [
        p.text.strip()
        for p in doc.paragraphs[section_index + 1 : section_index + 4]
    ]
    if clarification not in following_texts:
        # python-docx has no public insert-after helper; use XML insertion and then
        # reopen via the document paragraph list.
        from copy import deepcopy

        new_xml = deepcopy(section_paragraph._p)
        for child in list(new_xml):
            new_xml.remove(child)
        section_paragraph._p.addnext(new_xml)
        inserted = None
        for paragraph in doc.paragraphs:
            if paragraph._p is new_xml:
                inserted = paragraph
                break
        if inserted is None:
            raise RuntimeError("No se pudo insertar el párrafo aclaratorio")
        inserted.style = doc.paragraphs[section_index + 1].style
        inserted.add_run(clarification)
        set_run_style(inserted)

    target_table = None
    for table in doc.tables:
        if table.rows[0].cells[0].text.strip() == "Relación":
            target_table = table
            break

    if target_table is None:
        raise RuntimeError("No se encontró la tabla de relaciones")

    replacements = {
        "Relación": "Vínculo lógico",
        "Cardinalidad": "Tipo de vínculo",
        "Alerta - Reporte": "Alerta - Reporte",
        "1 a N": "Referencia 1 a N",
        "Cada reporte guarda alertaId como referencia al id de la alerta.": (
            "Cada reporte guarda alertaId como referencia lógica al id de la alerta. "
            "MongoDB no impone una clave foránea; la aplicación valida y consulta por ese campo."
        ),
        "N a N derivada": "Asociación derivada por consulta geoespacial",
        "No se guarda una tabla intermedia; se calcula por cercanía geográfica al momento de crear la alerta y se registra vecinosNotificados.": (
            "No se guarda una tabla intermedia. Se calcula por cercanía geográfica al momento de crear la alerta "
            "y se registra solamente vecinosNotificados como dato agregado."
        ),
        "1 a 1 opcional": "Sincronización 1 a 1 opcional",
        "Si el reporte tiene ubicación, se crea un nodo en Neo4j con mongoId igual al id del reporte MongoDB.": (
            "Si el reporte tiene ubicación, se crea un nodo en Neo4j con mongoId igual al id del reporte MongoDB "
            "para mantener trazabilidad entre ambos modelos."
        ),
        "1 a 1 secuencial": "Relación de grafo secuencial",
        "Cada nodo puede apuntar al siguiente avistamiento cronológico mediante SIGUIENTE.": (
            "En Neo4j sí existe una relación explícita: cada nodo puede apuntar al siguiente avistamiento "
            "cronológico mediante SIGUIENTE."
        ),
    }

    for row in target_table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if text in replacements:
                cell.text = replacements[text]
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(11)
                        if row == target_table.rows[0]:
                            run.bold = True

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
